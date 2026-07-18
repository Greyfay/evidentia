"""Load a heterogeneous dossier into provenance-bearing typed source tables.

Every native format (GDPdU index.xml + TXT, CSV, XLSX, DOCX, PDF) is parsed into a
uniform :class:`SourceTable`. Each cell can produce an immutable :class:`EvidenceRef`
bound to its exact source coordinate and raw value. No filename, vendor id, amount, or
label is ever assumed: tables are identified downstream by their columns, not their names.
"""

from __future__ import annotations

import codecs
import csv
import io
from dataclasses import dataclass
from pathlib import Path

from audit_compiler.adapters.gdpdu import compile_gdpdu_dossier
from audit_compiler.adapters.xlsx import parse_xlsx_workbook
from audit_compiler.inventory import inventory_dossier, sha256_file
from audit_compiler.models import DataLocale, EvidenceRef, SourceType


@dataclass(frozen=True)
class SourceTable:
    """A parsed source rendered as rows of ``column -> raw value`` with provenance."""

    name: str
    source_path: str
    file_sha256: str
    source_type: SourceType
    columns: tuple[str, ...]
    rows: tuple[dict[str, str], ...]
    row_numbers: tuple[int, ...]
    sheet: str | None = None
    page_numbers: tuple[int, ...] = ()

    def evidence(self, index: int, column: str, *, normalized: str | None = None) -> EvidenceRef:
        """Return an immutable evidence pointer for one cell of one row."""

        raw_value = self.rows[index].get(column, "")
        cell = None
        if self.source_type is SourceType.XLSX_CELL:
            cell = f"{_column_letter(self.columns.index(column))}{self.row_numbers[index]}"
        page = self.page_numbers[index] if self.page_numbers else None
        passage = raw_value[:200] if self.source_type is SourceType.PDF_PASSAGE else None
        return EvidenceRef.canonical(
            source_path=self.source_path,
            source_type=self.source_type,
            file_sha256=self.file_sha256,
            raw_value=raw_value,
            normalized_value=normalized,
            row=self.row_numbers[index] if self.source_type is not SourceType.PDF_PASSAGE else None,
            sheet=self.sheet,
            cell=cell,
            page=page,
            passage=passage,
        )


def _column_letter(zero_based: int) -> str:
    letters = ""
    n = zero_based + 1
    while n:
        n, remainder = divmod(n - 1, 26)
        letters = chr(65 + remainder) + letters
    return letters


def _decode(raw: bytes) -> str:
    if raw.startswith(codecs.BOM_UTF8):
        return raw.decode("utf-8-sig")
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        return raw.decode("cp1252")


def _rel(path: Path, root: Path) -> str:
    return path.resolve().relative_to(root.resolve()).as_posix()


def _load_header_csv(path: Path, root: Path) -> SourceTable:
    raw = path.read_bytes()
    text = _decode(raw)
    sample = text[:4096]
    delimiter = ";" if sample.count(";") >= sample.count(",") else ","
    reader = list(csv.reader(io.StringIO(text, newline=""), delimiter=delimiter))
    reader = [r for r in reader if any(cell.strip() for cell in r)]
    if not reader:
        raise ValueError(f"empty CSV: {path}")
    header = [h.strip() for h in reader[0]]
    rows: list[dict[str, str]] = []
    row_numbers: list[int] = []
    for physical_line, values in enumerate(reader[1:], start=2):
        padded = list(values) + [""] * (len(header) - len(values))
        rows.append({header[i]: padded[i] for i in range(len(header))})
        row_numbers.append(physical_line)
    return SourceTable(
        name=path.stem,
        source_path=_rel(path, root),
        file_sha256=sha256_file(path),
        source_type=SourceType.CSV_ROW,
        columns=tuple(header),
        rows=tuple(rows),
        row_numbers=tuple(row_numbers),
    )


def _load_xlsx(
    path: Path, root: Path, *, locale: DataLocale
) -> list[SourceTable]:
    """Adapt the one native XLSX parse into AIR source tables."""

    workbook = parse_xlsx_workbook(
        path, dossier_root=root, locale=locale.value
    )
    tables: list[SourceTable] = []
    for sheet in workbook.sheets:
        if sheet.header_row is None:
            continue
        cells = {(cell.row_number, cell.column_number): cell for cell in sheet.cells}
        rows: list[dict[str, str]] = []
        row_numbers: list[int] = []
        for record in sheet.records:
            values = {
                header: cells[(record.row_number, column)].raw_value
                for column, header in enumerate(sheet.headers, start=1)
            }
            if not any(values.values()):
                continue
            rows.append(values)
            row_numbers.append(record.row_number)
        tables.append(
            SourceTable(
                name=f"{path.stem}:{sheet.name}",
                source_path=workbook.source_file.path,
                file_sha256=workbook.source_file.sha256,
                source_type=SourceType.XLSX_CELL,
                columns=sheet.headers,
                rows=tuple(rows),
                row_numbers=tuple(row_numbers),
                sheet=sheet.name,
            )
        )
    return tables


def _load_docx(path: Path, root: Path) -> SourceTable:
    from docx import Document

    document = Document(str(path))
    rows: list[dict[str, str]] = []
    row_numbers: list[int] = []
    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text.strip()
        if text:
            rows.append({"paragraph": text})
            row_numbers.append(index)
    return SourceTable(
        name=path.stem,
        source_path=_rel(path, root),
        file_sha256=sha256_file(path),
        source_type=SourceType.DOCX_PARAGRAPH,
        columns=("paragraph",),
        rows=tuple(rows),
        row_numbers=tuple(row_numbers),
    )


def _load_pdf(path: Path, root: Path) -> SourceTable:
    import fitz  # PyMuPDF

    rows: list[dict[str, str]] = []
    page_numbers: list[int] = []
    with fitz.open(str(path)) as document:
        for page_index in range(document.page_count):
            text = document.load_page(page_index).get_text("text").strip()
            if text:
                rows.append({"text": text})
                page_numbers.append(page_index + 1)
    return SourceTable(
        name=path.stem,
        source_path=_rel(path, root),
        file_sha256=sha256_file(path),
        source_type=SourceType.PDF_PASSAGE,
        columns=("text",),
        rows=tuple(rows),
        row_numbers=tuple(range(1, len(rows) + 1)),
        page_numbers=tuple(page_numbers),
    )


@dataclass(frozen=True)
class LoadedDossier:
    """Every parsed source table plus the raw file manifest and per-file warnings."""

    root: Path
    locale: DataLocale
    tables: tuple[SourceTable, ...]
    warnings: tuple[tuple[str, str], ...]  # (source_path, message)

    def by_columns(self, required: set[str]) -> list[SourceTable]:
        """Return tables whose columns are a superset of ``required`` (exact header names)."""

        return [t for t in self.tables if required.issubset(set(t.columns))]


def load_dossier(
    directory: Path, *, locale: DataLocale | str = DataLocale.DE
) -> LoadedDossier:
    """Parse every supported file in a dossier into provenance-bearing source tables."""

    root = directory.expanduser().resolve()
    explicit_locale = DataLocale(locale)
    manifest = inventory_dossier(root)
    tables: list[SourceTable] = []
    warnings: list[tuple[str, str]] = []

    # GDPdU folders: parse each index.xml exactly once (its declared TXT/CSV files).
    gdpdu_data_paths: set[str] = set()
    for index_path in sorted(root.rglob("index.xml")):
        folder = index_path.parent
        folder_prefix = _rel(folder, root)
        folder_prefix = "" if folder_prefix == "." else f"{folder_prefix}/"
        try:
            # GDPdU data-file URLs are relative to the index.xml directory.
            compilation = compile_gdpdu_dossier(index_path, dossier_root=folder)
        except Exception as exc:  # noqa: BLE001 - surfaced as a warning, never invented data
            warnings.append((_rel(index_path, root), f"GDPdU compile failed: {exc}"))
            continue
        for parsed in compilation.tables:
            source_path = f"{folder_prefix}{parsed.definition.source_path}"
            gdpdu_data_paths.add(source_path)
            columns = tuple(c.name for c in parsed.definition.columns)
            source_type = (
                SourceType.CSV_ROW
                if parsed.source_file.file_type == "csv"
                else SourceType.TEXT_ROW
            )
            tables.append(
                SourceTable(
                    name=parsed.definition.table_name,
                    source_path=source_path,
                    file_sha256=parsed.source_file.sha256,
                    source_type=source_type,
                    columns=columns,
                    rows=tuple(record.named_values for record in parsed.records),
                    row_numbers=tuple(record.row_number for record in parsed.records),
                )
            )

    loaders = {
        "csv": _load_header_csv,
        "text": None,
        "xlsx": _load_xlsx,
        "docx": _load_docx,
        "pdf": _load_pdf,
    }
    for source in manifest.files:
        if source.path in gdpdu_data_paths or source.file_type in {"xml", "unknown", "xls"}:
            continue
        loader = loaders.get(source.file_type)
        if loader is None:
            continue
        path = root / source.path
        try:
            result = (
                _load_xlsx(path, root, locale=explicit_locale)
                if source.file_type == "xlsx"
                else loader(path, root)
            )
        except Exception as exc:  # noqa: BLE001 - never fabricate; record and continue
            warnings.append((source.path, f"parse failed: {exc}"))
            continue
        tables.extend(result if isinstance(result, list) else [result])

    return LoadedDossier(
        root=root,
        locale=explicit_locale,
        tables=tuple(tables),
        warnings=tuple(warnings),
    )
